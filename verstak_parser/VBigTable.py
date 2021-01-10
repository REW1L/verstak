from docx.table import Table, _Row, _Cell

from .VParagraph import VParagraph
from .VBoldText import VBoldText


class VBigTable:
    MAX_WIDTH = 700

    def __init__(self, table: Table = None, title: str = ""):
        self.title = title
        self.raw: Table = table
        self.headers = []
        self.rows = []
        if table is not None:
            self.parse(self.raw)

    def __str__(self):
        return "## BIG TABLE"

    def __check_headers(self):
        """
        Check the first row of the table for headers
        :return: True if headers found and False otherwise
        """
        headers = True
        for cell in self.raw.rows[0].cells:
            for paragraph in cell.paragraphs:
                vp = VParagraph(paragraph)
                for part in vp.parts:
                    if type(part) != VBoldText:
                        headers = False
                        break
                if not headers:
                    break
            if not headers:
                break
        return headers

    def __parse_headers(self, row: _Row):
        """
        Fill headers with contents of docx.table._Row
        :param row: array of docx.table._Row
        """
        for cell in row.cells:
            parts = []
            for paragraph in cell.paragraphs:
                parts.append(VParagraph(paragraph))
            self.headers.append(parts)

    def __parse_rows(self, rows: list):
        """
        Fill rows with contents of docx.table._Row array
        :param rows: array of docx.table._Row
        """
        for row in rows:
            row: _Row
            row_parts = []
            for cell in row.cells:
                cell: _Cell
                for paragraph in cell.paragraphs:
                    row_parts.append(VParagraph(paragraph))
            self.rows.append(row_parts)

    def parse(self, table: Table):
        """
        Parse table to get headers and rows
        """
        if len(table.rows) < 1:
            return
        elif len(table.rows) == 1:
            rows = table.rows
        elif not self.__check_headers():
            rows = table.rows
        else:
            self.__parse_headers(table.rows[0])
            rows = table.rows[1:]
        self.__parse_rows(rows)

    def to_html(self, skip: bool = False):
        """
        Get html representation
        :param skip: get stub instead of table
        :return: <p> with text BIG TABLE
        """
        if skip:
            return f"<p>## BIG TABLE</p>"
        html_table = ['<table class="desktop-table desktop-table--thead-with-border" style="width: '
                      f'{self.MAX_WIDTH}!important;">']
        if len(self.headers) != 0:
            col_width = int((self.MAX_WIDTH / len(self.headers)))
            html_table.append('<thead>')
            for header in self.headers:
                cell_html = [p.to_html() for p in header]
                html_table.append(f'<th style="width: {col_width}">{"".join(cell_html)}</th>')
            html_table.append('</thead>')
        html_table.append('<tbody>')
        for row in self.rows:
            html_table.append('<tr>')
            for cell in row:
                cell_html = [p.to_html() for p in cell]
                html_table.append(f'<td>{"".join(cell_html)}</td>')
            html_table.append('</tr>')
        html_table.append('</tbody>')
        html_table.append('</table>')
        return "\n".join(html_table)

    def do_typograf(self):
        pass
