import os
import re
import json
import docx
from lxml.etree import _Element
from docx.document import Document
from docx.oxml.text.run import CT_R
from docx.oxml.text.font import CT_RPr
from docx.oxml.text.parfmt import CT_PPr
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.numbering import CT_NumPr
from docx.text.paragraph import Paragraph
from docx.table import Table, _Row, _Cell, _Column
from docx.opc.rel import _Relationship

# Diego Perini's regex pattern for url
url_regex = "(?:(?:(?:https?|ftp):)?\\/\\/)(?:\\S+(?::\\S*)?@)?(?:(?!(?:10|1" \
            "27)(?:\\.\\d{1,3}){3})(?!(?:169\\.254|192\\.168)(?:\\.\\d{1,3})" \
            "{2})(?!172\\.(?:1[6-9]|2\\d|3[0-1])(?:\\.\\d{1,3}){2})(?:[1-9]" \
            "\\d?|1\\d\\d|2[01]\\d|22[0-3])(?:\\.(?:1?\\d{1,2}|2[0-4]\\d|25" \
            "[0-5])){2}(?:\\.(?:[1-9]\\d?|1\\d\\d|2[0-4]\\d|25[0-4]))|(?:(?:" \
            "[a-z0-9\\u00a1-\\uffff][a-z0-9\\u00a1-\\uffff_-]{0,62})?[a-z0-9" \
            "\\u00a1-\\uffff]\\.)+(?:[a-z\\u00a1-\\uffff]{2,}\\.?))(?::\\d{2" \
            ",5})?(?:[/?#]\\S*)?"


def parse_hyperlink(hyperlink: _Element, paragraph: Paragraph) -> str:
    url_id = None
    if len(hyperlink.values()) > 0:
        url_id = hyperlink.values()[0]
    url: _Relationship = paragraph.part.rels[url_id]
    url = url.target_ref
    text = hyperlink[0].text
    return f"[{text}]( {url} )"


def parse_bold(bold_text: CT_R) -> str:
    if bold_text.rPr.b is None:
        return f"{bold_text.text}"
    if bold_text.text.strip() == "":
        return bold_text.text
    bold_enabled = bold_text.rPr.b.val
    if bold_enabled:
        return f"**{bold_text.text}**"
    else:
        return f"{bold_text.text}"


def parse_list(list_text: CT_PPr, paragraph: Paragraph) -> str:
    option: CT_NumPr = list_text.numPr
    text = "  " * option.ilvl.val
    if paragraph.text.strip()[0].isupper() and paragraph.text.strip().endswith("."):
        text += "1. "
    else:
        text += "- "
    return text


def mock_picture(caption: str = "") -> str:
    return f'\n[img src="" prop="" caption="{caption}"]'


def parse_run(run: CT_R) -> str:
    line = ""
    if type(run[0]) == CT_RPr:
        if run[0].b is not None:
            line = parse_bold(run)
        elif len([tag for tag in run if tag.tag.endswith("drawing")]) > 0:
            line = mock_picture()
        else:
            line = run.text
    elif len([x for x in run if x.tag.endswith("drawing")]) > 0:
        line = mock_picture()
    return line


def parse_title(paragraph: Paragraph) -> list:
    text = []
    if paragraph.style.name.lower().startswith("title"):
        text.append("# ")
    elif paragraph.style.name.lower().startswith("heading "):
        heading_regexp = re.search("heading ([0-9]+)", paragraph.style.name.lower())
        if heading_regexp is not None and len(heading_regexp.groups()) > 0:
            heading_level = "#" * int(heading_regexp.group(1))
            text.append(f"#{heading_level} ")
    return text


def parse_paragraph(paragraph: Paragraph, title: bool = True) -> str:
    text = []
    if title:
        text = parse_title(paragraph)
    for elem in paragraph._element:
        line = ""
        if elem.tag.endswith("hyperlink"):
            line = parse_hyperlink(elem, paragraph)
            new_link = re.search(f"^\\[([^]]+)]\\( ({url_regex}) \\)", line)
            old_link = None
            if len(text) > 0:
                old_link = re.search(f"\\[([^]]+)]\\( ({url_regex}) \\)", text[-1])
                if old_link is not None:
                    if not new_link.group(1).startswith(old_link.group(1)):
                        text[-1] = "[{}]( {} )".format(old_link.group(1) + new_link.group(1), old_link.group(2))
                    else:
                        text[-1] = line
                    continue
        elif type(elem) == CT_PPr:
            if elem.numPr is not None:
                line = parse_list(elem, paragraph)
        elif type(elem) == CT_R:
            line = parse_run(elem)
            if len(text) > 0 and line.startswith("**") and text[-1].endswith("**"):
                line = line[2:]
                text[-1] = text[-1][:-2]
        else:
            if elem.text is not None:
                line = elem.text
        if line != "":
            text.append(line)
    return "".join(text)


def parse_pole(table: Table) -> list:
    section = []
    pole = {}
    for row in table.rows:
        left: _Cell = row.cells[0]
        right: _Cell = row.cells[1]
        pole["left"] = []
        pole["right"] = []
        section.append("```json")
        for para in left.paragraphs:
            text = parse_paragraph(para, False)
            if text.strip() != "":
                pole["left"].append(text)
        for para in right.paragraphs:
            text = parse_paragraph(para, False)
            if text.strip() != "":
                pole["right"].append(text)
        if len(pole["right"]) == 1:
            url_regexp = re.search(f"^\\[([^]]+)]\\( ({url_regex}) \\)$", pole["right"][0])
            if url_regexp is not None and len(url_regexp.groups()) == 2:
                pole["url"] = url_regexp.group(2)
                pole["right"] = url_regexp.group(1)
        for line in json.dumps(pole, sort_keys=False, indent=2, ensure_ascii=False).splitlines():
            section.append(line)
        section.append("```")
    return section


def parse_plashka(cell: _Cell) -> list:
    plashka = list()
    plashka.append("```plashka")
    for elem in cell._element:
        if type(elem) == CT_P:
            plashka.append(parse_paragraph(Paragraph(elem, cell)))
        elif type(elem) == CT_Tbl:
            plashka.extend(parse_table(Table(elem, cell)))
    plashka.append("```")
    return plashka


def parse_table(table: Table) -> list:
    section = []
    rows_len = len(table.rows)
    columns_len = len(table.columns)
    if rows_len < 1:
        return []
    if columns_len == 2:
        left_column: _Column = table.columns[0]
        right_column: _Column = table.columns[1]
        if float(left_column._gridCol.values()[0]) * 0.6 >= float(right_column._gridCol.values()[0]) and \
           len(left_column.cells[0].text) * 0.5 >= len(right_column.cells[0].text):
            section.extend(parse_pole(table))
        else:
            return ["## BIG TABLE"]
    elif columns_len == 1:
        section.extend(parse_plashka(table.rows[0].cells[0]))
    else:
        return ["## BIG TABLE"]
    if rows_len > 1 and columns_len == 1:
        rows = table.rows[1:]
        for row in rows:
            for cell in row.cells:
                section.append(cell.text)
    return section


def parse_document(document: Document):
    first_table_found = False
    first_title_found = False
    blank_line_caption_found = False
    text = []
    caption = False
    for elem in document.element.body:
        if type(elem) == CT_P:
            paragraph_text = parse_paragraph(Paragraph(elem, document))
            if not first_title_found and paragraph_text.startswith("# "):
                first_title_found = True
                continue
            if not caption or paragraph_text.find('img src="" prop="" caption') != -1 or \
               paragraph_text.startswith("#"):
                text.append(paragraph_text)
            else:
                if not blank_line_caption_found and paragraph_text.strip() == "":
                    blank_line_caption_found = True
                    continue
                imgs_count = 0
                for line in reversed(text):
                    if line.find('img src="" prop="" caption') == -1:
                        break
                    imgs_count += 1
                text[-imgs_count] = text[-imgs_count].replace('caption=""', f'caption="{paragraph_text.strip()}"', 1)
                caption = False
            if paragraph_text.find(mock_picture()) != -1:
                caption = True
        elif type(elem) == CT_Tbl:
            if not first_table_found:
                first_table_found = True
                continue
            text.extend(parse_table(Table(elem, document)))
            if len(text) > 2 and text[-1].find("BIG TABLE") != -1:
                text[-2] = f"## {text[-2]}"
        blank_line_caption_found = False
    return text


def parse_doc_file(doc_path: str) -> list:
    document: Document = docx.Document(doc_path)
    return parse_document(document)


def store_markdown(lines: list, path: str = f"markdown{os.sep}result.md"):
    with open(path, "w") as result_file:
        for line in lines:
            result_file.write(f"{line}\n")
        result_file.flush()


def main():
    docs = os.listdir("docs")
    i = 1
    for doc in docs:
        if not doc.startswith(".") and not doc.startswith("~"):
            print(f"Started parsing  ({i}/{len(docs)}): {doc}")
            lines = parse_doc_file(f"docs{os.sep}{doc}")
            store_markdown(lines, f"markdown{os.sep}{doc}.md")
            print(f"Finished parsing ({i}/{len(docs)}): {doc}")
        i += 1
    print("Done")


if __name__ == "__main__":
    main()
