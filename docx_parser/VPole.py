import json
from docx.table import Table, _Row, _Cell

from .VParagraph import VParagraph
from .VText import VText


class VPole:
    def __init__(self, row: _Row = None):
        """
        Class for representation of Pole
        :param row: row of the table with pole
        """
        self.title = ""
        self.left_parts = []
        self.right_parts = []
        self.url = ""
        self.raw = row
        self.removed_parts = []
        if row is not None:
            self.parse(row)

    def __dict__(self):
        d = {
            "left": [str(x) for x in self.left_parts],
            "right": [str(x) for x in self.right_parts]
        }
        if self.url != "":
            d["url"] = self.url
        return d

    def __str__(self):
        section = ["```json"]
        for line in json.dumps(self.__dict__(),
                               sort_keys=False,
                               indent=2,
                               ensure_ascii=False).splitlines():
            section.append(line)
        section.append("```")
        return "\n".join(section)

    def to_html(self):
        """
        Get html representation
        :return: html representation
        """
        url = ""
        if self.url != "":
            # links to journal.tinkoff.ru should be cut
            if self.url.startswith("https://journal.tinkoff.ru"):
                url = f' ref="{self.url[27:]}"'
            else:
                url = f' url="{self.url}"'
        html_parts = [f'<div class="with-aside">', self.left_parts[0].to_html(), f"[aside{url}]"]
        if len(self.right_parts) > 0:
            html_parts.extend([x.to_html() for x in self.right_parts])
            html_parts.append("[/aside]")
        html_parts.append("</div>")
        # all of the paragraphs after the first one on the left part should go after pole
        for part in self.left_parts[1:]:
            html_parts.append(part.to_html())
        return "\n".join(html_parts)

    def __clean_right_parts(self):
        """
        Clean right part from urls if there are some links to journal.tinkoff.ru
        """
        parts = []
        if len(self.right_parts) > 0:
            if self.url.startswith("https://journal.tinkoff.ru"):
                self.removed_parts = self.right_parts
                self.right_parts = []
                return
            remove_links = 0
            for paragraph in self.right_parts:
                if str(paragraph).find("https://journal.tinkoff.ru") == -1:
                    remove_links |= 1
                else:
                    remove_links |= 2
            if remove_links == 3:
                for part in self.right_parts:
                    if str(part).find("https://journal.tinkoff.ru") == -1:
                        parts.append(part)
                    else:
                        self.removed_parts.append(part)
            else:
                parts = self.right_parts
            if len(parts) == 1 and type(parts[0]) == VParagraph:
                links = parts[0].get_links_indexes()
                if len(links) == 1 and parts[0].text.strip() == parts[0][links[0]].text.strip():
                    self.url = parts[0][links[0]].url
                    parts[0] = VText(parts[0].text)
        self.right_parts = parts

    def parse(self, row: _Row):
        """
        Parse row to VPole
        :param row: row from docx.table.Table
        """
        left: _Cell = row.cells[0]
        right: _Cell = row.cells[1]
        self.left_parts = []
        self.right_parts = []
        for para in left.paragraphs:
            paragraph = VParagraph(para, False)
            if str(paragraph).strip() != "":
                self.left_parts.append(paragraph)
        for para in right.paragraphs:
            paragraph = VParagraph(para, False)
            if str(paragraph).strip() != "":
                self.right_parts.append(paragraph)
        if len(self.right_parts) == 1:
            links = self.right_parts[0].get_links_indexes()
            if len(links) == 1 and self.right_parts[0].text.strip() == self.right_parts[0][links[0]].text.strip():
                self.url = self.right_parts[0][links[0]].url
                self.right_parts[0] = VText(self.right_parts[0].text)
        self.__clean_right_parts()

    def do_typograf(self):
        """
        Rework text of parts by rules from typograph
        :return: resulted text
        """
        self.left_parts[0].do_typograf(nobr_enabled=False)
        for paragraph in self.left_parts[1:]:
            paragraph.do_typograf()
        for paragraph in self.right_parts:
            paragraph.do_typograf(nobr_enabled=False)

    @staticmethod
    def parse_poles(table: Table) -> list:
        """
        Represent table as multiple poles and parse it
        :param table: docx.table.Table to parse
        :return: list of VPoles
        """
        poles = []
        for row in table.rows:
            right_text = "".join([VParagraph(x).text for x in row.cells[1].paragraphs])
            if right_text == "":
                for para in row.cells[0].paragraphs:
                    poles.append(VParagraph(para))
            else:
                poles.append(VPole(row))
        return poles
